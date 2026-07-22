"""
Resume text extraction service.

Handles PDF and DOCX file parsing with detection of tricky edge cases:
- Scanned/image-only PDFs (no extractable text)
- Multi-column layouts (detected via bounding-box cluster analysis)

Architectural decision: We extract text+position data from PDFs using PyMuPDF,
but deliberately do NOT attempt to re-order multi-column content. The structurer
service receives the raw text and any positional metadata — it can then decide
how to handle it. This keeps the parser single-responsibility.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class ParseResult:
    """Result of the extraction step, before structuring."""
    def __init__(
        self,
        raw_text: str,
        file_type: str,
        is_scanned: bool = False,
        is_multi_column: bool = False,
        page_count: int = 0,
        warnings: Optional[list[str]] = None,
    ):
        self.raw_text = raw_text
        self.file_type = file_type
        self.is_scanned = is_scanned
        self.is_multi_column = is_multi_column
        self.page_count = page_count
        self.warnings = warnings or []

    @property
    def success(self) -> bool:
        return bool(self.raw_text.strip()) and not self.is_scanned


def extract_from_pdf(file_bytes: bytes) -> ParseResult:
    """Extract text from a PDF using PyMuPDF.

    Returns a ParseResult with raw text, metadata, and any warnings.
    Raises ValueError if PyMuPDF is not installed or file is corrupt.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF parsing: pip install PyMuPDF")

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        raise ValueError(f"Cannot open PDF file: {e}")

    page_count = len(doc)
    logger.info(f"PDF has {page_count} page(s)")

    # Check if first page has any text
    first_page = doc[0]
    first_blocks = first_page.get_text("blocks")
    total_text_length = sum(len(b[4]) for b in first_blocks)

    if total_text_length < 10 and page_count > 0:
        doc.close()
        return ParseResult(
            raw_text="",
            file_type="pdf",
            is_scanned=True,
            page_count=page_count,
            warnings=["This appears to be a scanned/image-based PDF."],
        )

    # Extract all text with position data for multi-column detection
    all_blocks: list[tuple[float, float, float, float, str]] = []  # (x0, y0, x1, y1, text)
    for page in doc:
        blocks = page.get_text("blocks")
        for b in blocks:
            # PyMuPDF block: (x0, y0, x1, y1, text, block_no, block_type)
            if b[6] == 0:  # text block (not image)
                x0, y0, x1, y1 = b[0], b[1], b[2], b[3]
                text = b[4].strip()
                if text:
                    all_blocks.append((x0, y0, x1, y1, text))

    doc.close()

    # Detect multi-column layout by analyzing x-coordinate clusters
    is_multi_column = _detect_multi_column(all_blocks)

    # Sort blocks by y-position (top-to-bottom), then x (left-to-right)
    # For multi-column docs this won't give perfect reading order,
    # but we flag it explicitly rather than trying complex reconstruction.
    all_blocks.sort(key=lambda b: (b[1], b[0]))  # (y0, x0)

    raw_text = "\n".join(b[4] for b in all_blocks)

    warnings: list[str] = []
    if is_multi_column:
        warnings.append(
            "Multi-column layout detected. Consider using a single-column "
            "format for better ATS parsing compatibility."
        )

    return ParseResult(
        raw_text=raw_text,
        file_type="pdf",
        is_scanned=False,
        is_multi_column=is_multi_column,
        page_count=page_count,
        warnings=warnings,
    )


def extract_from_docx(file_bytes: bytes) -> ParseResult:
    """Extract text from a DOCX file using python-docx.

    DOCX files are generally easier to parse than PDFs — they have
    native paragraph structure and no layout ambiguity.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing: pip install python-docx")

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Cannot open DOCX file: {e}")

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    raw_text = "\n".join(paragraphs)

    # Detect tables — these can cause issues for ATS parsers
    has_tables = len(doc.tables) > 0
    warnings: list[str] = []
    if has_tables:
        warnings.append(
            "Tables detected in your resume. Some ATS parsers struggle with "
            "table-formatted content. Consider using a clean text layout."
        )

    return ParseResult(
        raw_text=raw_text,
        file_type="docx",
        is_scanned=False,
        is_multi_column=False,
        page_count=len(paragraphs) // 20 + 1,  # rough estimate
        warnings=warnings,
    )


def _detect_multi_column(
    blocks: list[tuple[float, float, float, float, str]],
    threshold: float = 0.3,
) -> bool:
    """Heuristic: detect multi-column layout from text block x-coordinates.

    If a significant portion of text blocks occupy distinct horizontal ranges
    (i.e., there's a clear vertical gap running through the page), it's likely
    multi-column.

    Args:
        blocks: List of (x0, y0, x1, y1, text) tuples.
        threshold: Fraction of blocks that must be in a secondary column.

    Returns:
        True if multi-column layout is detected.
    """
    if not blocks:
        return False

    # Get x-centers of all blocks
    x_centers = [(b[0] + b[2]) / 2 for b in blocks]
    if not x_centers:
        return False

    # Find the page width (max x1 - min x0)
    min_x = min(b[0] for b in blocks)
    max_x = max(b[2] for b in blocks)
    page_width = max_x - min_x
    if page_width == 0:
        return False

    # Normalize x-centers to 0-1 range
    normalized = [(x - min_x) / page_width for x in x_centers]

    # If there's a clear vertical gap (e.g., 0.35–0.65 with few blocks),
    # it's multi-column. Count blocks in the "gap" region.
    gap_region = [n for n in normalized if 0.35 < n < 0.65]
    gap_ratio = len(gap_region) / len(normalized)

    # Also check if blocks cluster into two groups (left/right)
    left_cluster = sum(1 for n in normalized if n < 0.35)
    right_cluster = sum(1 for n in normalized if n > 0.65)
    min_cluster = min(left_cluster, right_cluster)
    cluster_ratio = min_cluster / len(normalized) if len(normalized) > 0 else 0

    return gap_ratio < 0.15 and cluster_ratio > threshold


def extract_text(file_bytes: bytes, filename: str) -> ParseResult:
    """Dispatch to the appropriate parser based on file extension.

    Args:
        file_bytes: Raw file content.
        filename: Original filename for extension detection.

    Returns:
        ParseResult with extracted text and metadata.

    Raises:
        ValueError: If the file type is unsupported.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_from_pdf(file_bytes)
    elif ext == ".docx":
        return extract_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx")
